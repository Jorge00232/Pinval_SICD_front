from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path("Informe_Desarrollo_SICD_Pinval.docx")
BLUE = "2563EB"
NAVY = "172033"
GREEN = "16805B"
LIGHT_BLUE = "EAF2FF"
LIGHT_GRAY = "F3F6FA"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=NAVY, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        shade(table.rows[0].cells[index], BLUE)
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=WHITE, size=9)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if row_index % 2 == 1:
                shade(cells[index], LIGHT_GRAY)
            set_cell_text(cells[index], value, size=8.5)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def add_title(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else NAVY)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    return paragraph


def add_bullets(doc, items, level=0):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)


def add_code(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.right_indent = Cm(0.7)
    paragraph.paragraph_format.space_after = Pt(8)
    shade_paragraph = OxmlElement("w:shd")
    shade_paragraph.set(qn("w:fill"), "EEF2F7")
    paragraph._p.get_or_add_pPr().append(shade_paragraph)
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title + "\n")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    body = paragraph.add_run(text)
    body.font.name = "Aptos"
    body.font.size = Pt(9)
    body.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(34)
styles["Title"].font.bold = True
styles["Heading 1"].font.name = "Aptos Display"
styles["Heading 1"].font.size = Pt(20)
styles["Heading 1"].font.bold = True
styles["Heading 2"].font.name = "Aptos Display"
styles["Heading 2"].font.size = Pt(15)
styles["Heading 2"].font.bold = True
styles["Heading 3"].font.name = "Aptos"
styles["Heading 3"].font.size = Pt(12)
styles["Heading 3"].font.bold = True

# Header and footer
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header.add_run("SICD PINVAL | INFORME DE DESARROLLO")
run.font.name = "Aptos"
run.font.size = Pt(8)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(BLUE)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("Documento técnico de construcción y evolución del sistema")
footer_run.font.name = "Aptos"
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor.from_string("64748B")

# Cover
for _ in range(4):
    doc.add_paragraph()
badge = doc.add_paragraph()
badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
badge_run = badge.add_run("PROYECTO DE MODERNIZACIÓN")
badge_run.bold = True
badge_run.font.name = "Aptos"
badge_run.font.size = Pt(11)
badge_run.font.color.rgb = RGBColor.from_string(BLUE)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Informe de construcción y desarrollo del SICD Pinval")
title_run.font.color.rgb = RGBColor.from_string(NAVY)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run(
    "Sistema interno para control de inventario, productos, compras, ventas, trazabilidad y consultas asistidas"
)
subtitle_run.font.name = "Aptos"
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = RGBColor.from_string("64748B")

doc.add_paragraph()
add_callout(
    doc,
    "Objetivo del documento",
    "Registrar paso a paso la evolución funcional y técnica del proyecto, las decisiones tomadas, "
    "las validaciones realizadas y los próximos trabajos necesarios para convertir el prototipo en "
    "una solución operativa conectada a datos reales.",
)

cover_meta = doc.add_paragraph()
cover_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = cover_meta.add_run("Fecha de elaboración: 7 de junio de 2026\nEstado: prototipo funcional en desarrollo")
meta.font.name = "Aptos"
meta.font.size = Pt(10)
meta.font.color.rgb = RGBColor.from_string("64748B")

add_page_break(doc)

add_title(doc, "Contenido", 1)
contents = [
    "1. Resumen ejecutivo",
    "2. Contexto, problema y objetivos",
    "3. Alcance y arquitectura actual",
    "4. Construcción del frontend paso a paso",
    "5. Construcción del backend paso a paso",
    "6. Tratamiento y normalización de datos",
    "7. Integración del chatbot",
    "8. Seguridad, autenticación, roles y auditoría",
    "9. Pruebas y validaciones realizadas",
    "10. Ejecución local y dependencias",
    "11. Decisiones técnicas y justificación",
    "12. Limitaciones, riesgos y trabajo pendiente",
    "13. Próximos pasos recomendados",
    "14. Conclusiones",
]
add_numbered(doc, contents)

add_page_break(doc)

add_title(doc, "1. Resumen ejecutivo", 1)
add_body(
    doc,
    "SICD Pinval es un sistema web interno orientado a modernizar el control de inventario de Pinval. "
    "El proyecto nació a partir de una operación gestionada principalmente mediante planillas Excel y "
    "un sistema POS que no se encontraba conectado al inventario. Esta separación producía información "
    "desactualizada, mayor riesgo de errores manuales y poca trazabilidad sobre entradas, salidas y ajustes."
)
add_body(
    doc,
    "Durante el desarrollo se construyó una interfaz web con vistas especializadas para dashboard, ventas, "
    "compras, inventario, productos, movimientos, clientes, proveedores, alertas y reportes. Posteriormente "
    "se conectó el frontend con un backend NestJS que consulta datos reales almacenados en SQLite mediante Prisma."
)
add_body(
    doc,
    "Como evolución adicional se incorporó un chatbot interno de solo lectura. Este asistente permite consultar "
    "stock por código o nombre, productos bajo mínimo, productos sin stock, productos con stock negativo, "
    "rankings de mayor o menor stock y resúmenes generales del inventario. Finalmente, se protegió el chatbot "
    "mediante autenticación JWT, roles y auditoría sanitizada."
)
add_callout(
    doc,
    "Estado actual",
    "El sistema funciona como prototipo integrado. El frontend y backend compilan correctamente; el chatbot "
    "consulta datos reales; las consultas protegidas exigen JWT; y las auditorías quedan almacenadas en SQLite. "
    "Aún faltan endpoints transaccionales completos, usuarios persistidos en base de datos y despliegue productivo.",
    fill="E8F7EF",
)

add_title(doc, "2. Contexto, problema y objetivos", 1)
add_title(doc, "2.1 Contexto operacional", 2)
add_bullets(
    doc,
    [
        "Pinval comercializa productos y opera principalmente en un contexto B2B, manteniendo también atención B2C.",
        "Las ventas se registran en un sistema POS, pero este no se encuentra integrado directamente con el inventario.",
        "El inventario y parte de la información comercial provenían de archivos Excel.",
        "Los archivos de origen contienen datos históricos, diferencias de formato y registros con stock negativo.",
        "El proyecto separa el trabajo de frontend y backend para permitir desarrollo paralelo entre integrantes.",
    ],
)

add_title(doc, "2.2 Problemas identificados", 2)
add_table(
    doc,
    ["Problema", "Consecuencia operacional", "Respuesta del proyecto"],
    [
        ["Dependencia de Excel", "Actualización manual y riesgo de versiones distintas.", "Centralización progresiva mediante API y base de datos."],
        ["POS sin conexión al inventario", "Las ventas no descuentan automáticamente el stock real.", "Diseño de flujo de ventas y movimientos; integración completa pendiente."],
        ["Stock negativo en datos originales", "Indica inconsistencias que requieren revisión, no stock disponible.", "Se conserva el valor original y se clasifica como requiere ajuste."],
        ["Información excesiva en pantalla", "Dificulta comprender el estado operativo.", "Rediseño minimalista, filtros, resúmenes y vistas por función."],
        ["Falta de trazabilidad", "No se conoce claramente quién consultó o modificó información.", "Movimientos, roles y auditoría de consultas del chatbot."],
    ],
    widths=[4, 6, 6],
)

add_title(doc, "2.3 Objetivo general", 2)
add_body(
    doc,
    "Desarrollar un sistema web centralizado que permita consultar y administrar información de inventario, "
    "productos, compras y ventas de Pinval, reduciendo errores manuales y mejorando la trazabilidad."
)

add_title(doc, "2.4 Objetivos específicos", 2)
add_bullets(
    doc,
    [
        "Organizar las vistas según procesos operativos comprensibles.",
        "Conectar el frontend con datos reales entregados por el backend.",
        "Normalizar códigos, stock, precios, familias y fechas provenientes de las fuentes originales.",
        "Diferenciar stock disponible, bajo mínimo, sin stock y registros que requieren ajuste.",
        "Permitir consultas rápidas mediante un chatbot interno de solo lectura.",
        "Aplicar autenticación, roles, protección de rutas y auditoría.",
        "Mantener una arquitectura que permita migrar posteriormente desde SQLite hacia PostgreSQL.",
    ],
)

add_page_break(doc)
add_title(doc, "3. Alcance y arquitectura actual", 1)
add_title(doc, "3.1 Componentes principales", 2)
add_table(
    doc,
    ["Capa", "Tecnologías", "Responsabilidad"],
    [
        ["Frontend", "React 19, TypeScript 6, Vite 8, React Router", "Interfaz, navegación, vistas operativas, sesión y chatbot flotante."],
        ["Backend", "NestJS 11, TypeScript, Passport, JWT, bcrypt", "API, reglas de consulta, autenticación, autorización y auditoría."],
        ["Acceso a datos", "Prisma 7, Prisma Adapter libSQL", "Consulta tipada y desacoplamiento de la base de datos."],
        ["Base de datos actual", "SQLite / libSQL", "Datos importados de stock, ventas y auditoría del chatbot."],
        ["Pruebas", "Jest, ESLint, TypeScript build", "Validación automatizada y control de calidad."],
    ],
    widths=[3, 6, 7],
)

add_title(doc, "3.2 Flujo general", 2)
add_code(
    doc,
    "Usuario -> Frontend React -> API NestJS -> Servicios de negocio -> Prisma -> SQLite\n"
    "                    |-> JWT y roles\n"
    "                    |-> Chatbot de consultas controladas\n"
    "                    |-> Auditoría sanitizada"
)
add_body(
    doc,
    "El frontend no debe acceder directamente a SQLite. Toda consulta o futura modificación debe pasar por el "
    "backend, donde se validan credenciales, permisos, tipos de datos y reglas de negocio."
)

add_title(doc, "3.3 Repositorios y copias de trabajo", 2)
add_body(
    doc,
    "Para evitar interferir con los cambios simultáneos del equipo, la integración reciente del chatbot y la "
    "seguridad se trabajaron en copias aisladas sin repositorio Git:"
)
add_code(
    doc,
    "Frontend aislado: C:\\Users\\const\\OneDrive\\Desktop\\Documentos\\Pinval_SICD_front-2\n"
    "Backend aislado:  C:\\Users\\const\\OneDrive\\Desktop\\Documentos\\Pinval_SICD_back-2"
)
add_callout(
    doc,
    "Consideración de integración",
    "Antes de incorporar estos cambios al repositorio principal se debe comparar el estado actual del equipo, "
    "resolver diferencias y realizar una integración controlada mediante rama o pull request.",
    fill="FFF4DB",
)

add_page_break(doc)
add_title(doc, "4. Construcción del frontend paso a paso", 1)
add_title(doc, "4.1 Eliminación de la carga de Excel desde la interfaz", 2)
add_body(
    doc,
    "Inicialmente existía una vista para cargar archivos Excel. Se eliminó porque el objetivo definitivo es que "
    "los datos provengan de una base de datos administrada por el backend. Mantener una carga manual en el frontend "
    "habría duplicado responsabilidades y creado una fuente adicional de errores."
)

add_title(doc, "4.2 Reorganización de navegación y vistas", 2)
add_body(
    doc,
    "La navegación se ordenó según un flujo operativo: operación, gestión y control. También se actualizaron las "
    "vistas disponibles después de eliminar módulos que ya no tenían sentido."
)
add_table(
    doc,
    ["Vista", "Propósito actual", "Aspectos trabajados"],
    [
        ["Dashboard", "Resumen operativo general.", "Indicadores, stock, alertas, actividad y diseño más compacto."],
        ["Ventas", "Registrar y consultar salidas.", "Cliente opcional, documento, productos, cantidades y permisos por rol."],
        ["Compras", "Registrar y consultar entradas.", "Proveedor, número de factura único y múltiples productos por compra."],
        ["Inventario", "Consultar stock y valorización.", "Estados claros, fechas, costos y reducción de información confusa."],
        ["Productos", "Mantener catálogo maestro.", "Campos alineados con CSV, categorías y diseño minimalista."],
        ["Movimientos", "Centralizar trazabilidad.", "Entradas, salidas, usuarios, fechas y detalles."],
        ["Clientes", "Mantener clientes recurrentes.", "Tratamiento B2B/B2C e identificadores."],
        ["Proveedores", "Mantener información de proveedores.", "Ampliación y orden de campos."],
        ["Alertas", "Priorizar productos que requieren atención.", "Separación entre bajo mínimo, sin stock y ajustes."],
        ["Reportes", "Seleccionar y generar consultas operativas.", "Tipos de reporte y textos más claros."],
    ],
    widths=[3, 5, 8],
)

add_title(doc, "4.3 Evolución visual", 2)
add_bullets(
    doc,
    [
        "Se reemplazó una interfaz inicialmente tosca por una composición más limpia y minimalista.",
        "Se redujeron tarjetas y textos innecesarios para priorizar información operacional.",
        "Se evitó el desplazamiento horizontal en tablas siempre que la información podía reorganizarse.",
        "Se aplicó una identidad personalizada para Pinval.",
        "Se incorporaron colores funcionales para distinguir operación, gestión, alertas y estados.",
        "Se trabajó el modo oscuro para mantener contraste y legibilidad.",
        "Se agregaron controles de idioma español e inglés mediante una carpeta centralizada de lenguaje.",
    ],
)

add_title(doc, "4.4 Roles visibles en la interfaz", 2)
add_table(
    doc,
    ["Rol", "Uso esperado", "Comportamiento del frontend"],
    [
        ["ADMIN", "Administración general.", "Acceso a consulta y gestión completa."],
        ["STOCK", "Control de inventario.", "Acceso a operaciones relacionadas con stock."],
        ["VIEWER", "Consulta.", "Debe visualizar datos sin formularios de modificación."],
    ],
    widths=[3, 6, 7],
)

add_title(doc, "4.5 Internacionalización", 2)
add_body(
    doc,
    "Se incorporó un selector ES/EN y se centralizaron textos traducibles. La solución se diseñó para evitar "
    "traducciones dispersas en cada componente. La traducción de términos del dominio debe mantenerse manual y "
    "controlada para no alterar conceptos de inventario."
)

add_title(doc, "4.6 Protección de rutas y sesión", 2)
add_bullets(
    doc,
    [
        "La sesión autenticada se almacena en sessionStorage.",
        "Las rutas internas utilizan ProtectedRoute.",
        "Un usuario sin sesión es redirigido al login.",
        "El cierre de sesión elimina el token.",
        "El chatbot envía el JWT mediante el encabezado Authorization.",
        "Los métodos de acceso simulados que generaban tokens falsos fueron deshabilitados.",
    ],
)

add_page_break(doc)
add_title(doc, "5. Construcción del backend paso a paso", 1)
add_title(doc, "5.1 Puesta en marcha inicial", 2)
add_body(
    doc,
    "El backend fue revisado y ejecutado como una aplicación NestJS. Inicialmente exponía rutas simples para stock "
    "y ventas. Se confirmó que compilaba en modo observación y que podía leer los datos importados."
)
add_code(
    doc,
    "npm install\n"
    "npx prisma generate\n"
    "npm run start:dev"
)

add_title(doc, "5.2 Normalización de productos", 2)
add_body(
    doc,
    "Se creó una respuesta normalizada de productos que combina datos de stock valorizado y ventas. Esta capa evita "
    "que el frontend dependa directamente de las columnas originales de Excel o SQLite."
)
add_table(
    doc,
    ["Campo normalizado", "Origen o regla", "Uso"],
    [
        ["codigo", "Código convertido a texto y completado a seis dígitos.", "Identificador consistente."],
        ["descrip", "Descripción de stock o ventas.", "Nombre mostrado al usuario."],
        ["familia", "Familia de ventas o valor por defecto.", "Clasificación."],
        ["stock", "Máximo entre stock original y cero.", "Stock disponible mostrado."],
        ["stockOriginal", "Valor real importado.", "Permite detectar negativos."],
        ["dataIssue", "STOCK_NEGATIVO cuando corresponde.", "Clasifica registros que requieren ajuste."],
        ["prcosto / prventa", "Precios disponibles.", "Valorizaciones y análisis."],
        ["minStock", "Actualmente valor temporal 5.", "Alertas de reposición."],
        ["fecha", "Fecha disponible en stock valorizado.", "Análisis temporal futuro."],
    ],
    widths=[3, 7, 6],
)

add_title(doc, "5.3 Endpoints disponibles", 2)
add_table(
    doc,
    ["Método", "Ruta", "Estado y propósito"],
    [
        ["GET", "/", "Estado básico de la API."],
        ["POST", "/auth/login", "Valida credenciales y entrega JWT."],
        ["GET", "/products", "Entrega catálogo normalizado."],
        ["GET", "/stock", "Consulta datos de stock."],
        ["GET", "/ventas", "Consulta datos importados de ventas."],
        ["GET", "/chatbot", "Informa capacidades del chatbot."],
        ["POST", "/chatbot/message", "Consulta protegida mediante JWT y roles."],
    ],
    widths=[2.5, 5, 8.5],
)

add_title(doc, "5.4 Preparación para migración de base de datos", 2)
add_body(
    doc,
    "SQLite se utiliza durante el desarrollo porque simplifica la ejecución local y permite trabajar con los datos "
    "importados. Prisma desacopla parte de la lógica para facilitar una futura migración a PostgreSQL. Sin embargo, "
    "el cambio no será completamente automático: será necesario ajustar el datasource, migraciones, tipos y pruebas."
)

add_page_break(doc)
add_title(doc, "6. Tratamiento y normalización de datos", 1)
add_title(doc, "6.1 Fuentes revisadas", 2)
add_bullets(
    doc,
    [
        "ventas.XLS y ventas.csv.",
        "stockvalorizado.xls y stockvalorizado.csv.",
        "Tablas SQLite equivalentes a ventas y stock valorizado.",
    ],
)
add_body(
    doc,
    "Los CSV extraídos desde SQLite permitieron comprobar que la base de datos mantenía la estructura original y "
    "que existían diferencias relevantes, como códigos con ceros iniciales, stock negativo, precios faltantes y "
    "familias sin definición."
)

add_title(doc, "6.2 Tratamiento del stock negativo", 2)
add_body(
    doc,
    "Se confirmó que los valores negativos provenían de la gestión original del negocio y no debían eliminarse. "
    "La solución adoptada evita mostrarlos como stock disponible, pero conserva stockOriginal y marca dataIssue "
    "como STOCK_NEGATIVO. De esta forma, el sistema distingue claramente entre un producto sin unidades y un dato "
    "que requiere revisión o regularización."
)
add_callout(
    doc,
    "Regla aplicada",
    "Stock disponible mostrado = máximo entre stock original y cero. "
    "Si stockOriginal es negativo, el registro se clasifica como REQUIERE_AJUSTE.",
)

add_title(doc, "6.3 Valorizaciones y fechas", 2)
add_body(
    doc,
    "Se detectó que mostrar el valor total potencial de todos los productos como si fuera ganancia podía inducir a "
    "error. La valorización de inventario debe diferenciar costo actual del inventario, valor aproximado de venta y "
    "ganancia efectivamente realizada dentro de un período. Para esto se requiere completar la trazabilidad temporal "
    "de compras y ventas; actualmente no toda esa información está disponible como transacción individual."
)

add_page_break(doc)
add_title(doc, "7. Integración del chatbot", 1)
add_title(doc, "7.1 Objetivo y alcance", 2)
add_body(
    doc,
    "El chatbot fue incorporado como un asistente interno flotante visible después de iniciar sesión. Su función es "
    "facilitar consultas de inventario usando lenguaje cotidiano, sin modificar directamente la base de datos."
)

add_title(doc, "7.2 Construcción inicial", 2)
add_numbered(
    doc,
    [
        "Se creó el módulo, controlador y servicio chatbot en NestJS.",
        "Se agregó GET /chatbot para consultar estado y capacidades.",
        "Se agregó POST /chatbot/message para recibir preguntas.",
        "Se implementó inicialmente la búsqueda de stock por código o nombre.",
        "Se construyó un botón flotante global en el frontend.",
        "Se agregó historial de mensajes, sugerencias, indicador de carga y manejo de errores.",
    ],
)

add_title(doc, "7.3 Consultas controladas implementadas", 2)
add_table(
    doc,
    ["Intención", "Ejemplos reconocidos", "Respuesta"],
    [
        ["Producto específico", "stock 005604; stock cloro", "Ficha del producto encontrado."],
        ["Resumen", "resumen de inventario", "Totales y estados generales."],
        ["Bajo mínimo", "productos bajo mínimo; por reponer", "Productos positivos iguales o inferiores al mínimo."],
        ["Sin stock", "productos sin stock; agotados", "Productos en cero, excluyendo negativos."],
        ["Requiere ajuste", "stock negativo; productos que requieren ajuste", "Registros con problema de datos."],
        ["Menor stock", "productos con menos stock; cuáles quedan menos", "Ranking ascendente de stock válido."],
        ["Mayor stock", "productos con mayor stock; más unidades", "Ranking descendente de stock válido."],
    ],
    widths=[3.5, 6.5, 6],
)

add_title(doc, "7.4 Decisión de no incorporar todavía un modelo generativo", 2)
add_body(
    doc,
    "Se decidió no integrar todavía un modelo de inteligencia artificial generativa. Primero se construyó un "
    "intérprete determinístico, porque las consultas principales pueden resolverse de forma segura, predecible y "
    "auditable. Un modelo futuro debería limitarse a identificar intenciones y nunca ejecutar SQL libre ni modificar "
    "datos directamente."
)
add_code(
    doc,
    "Pregunta libre -> Identificación de intención -> Consulta permitida del backend -> Datos reales -> Respuesta"
)

add_title(doc, "7.5 Evolución de la interfaz del chatbot", 2)
add_bullets(
    doc,
    [
        "Primera versión: tarjetas grandes por producto y scrollbar horizontal.",
        "Segunda versión: filas compactas con stock alineado y estados mediante color.",
        "Versión actual: muestra inicialmente tres resultados y permite desplegar los demás con Mostrar más.",
        "Se eliminó el desplazamiento horizontal interno.",
        "Se mantuvo compatibilidad con modo oscuro.",
        "Las listas se limitan desde backend a diez resultados para evitar saturación.",
    ],
)

add_page_break(doc)
add_title(doc, "8. Seguridad, autenticación, roles y auditoría", 1)
add_title(doc, "8.1 Problema de seguridad inicial", 2)
add_body(
    doc,
    "El login inicial devolvía un UUID aleatorio que no era validado posteriormente. Además, el frontend podía "
    "crear sesiones simuladas cuando el backend no estaba disponible. Esto permitía entrar a la interfaz con tokens "
    "falsos y el endpoint del chatbot podía consultarse sin autenticación."
)

add_title(doc, "8.2 JWT y contraseñas", 2)
add_bullets(
    doc,
    [
        "El backend ahora genera JWT firmados con duración de ocho horas.",
        "El token contiene identificador, usuario, nombre y rol.",
        "Las contraseñas de prueba se comparan contra hashes bcrypt.",
        "El secreto JWT se configura mediante JWT_SECRET.",
        "El frontend almacena la sesión en sessionStorage.",
        "Al cerrar la pestaña o sesión, el token deja de estar disponible.",
    ],
)

add_title(doc, "8.3 Autorización por roles", 2)
add_body(
    doc,
    "Se implementaron JwtAuthGuard, RolesGuard y el decorador Roles. El endpoint POST /chatbot/message acepta "
    "actualmente ADMIN, STOCK y VIEWER porque el chatbot es de solo lectura. Si en el futuro se incorporan acciones "
    "de modificación, estas deberán restringirse explícitamente."
)

add_title(doc, "8.4 Auditoría y sanitización", 2)
add_table(
    doc,
    ["Campo de auditoría", "Descripción"],
    [
        ["userId", "Identificador del usuario autenticado."],
        ["username", "Nombre de cuenta que realizó la consulta."],
        ["role", "Rol utilizado al consultar."],
        ["query", "Consulta sanitizada y limitada a 200 caracteres."],
        ["responseType", "Tipo de respuesta generada."],
        ["createdAt", "Fecha y hora automática."],
    ],
    widths=[4, 12],
)
add_body(
    doc,
    "Antes de almacenar la consulta, el backend reemplaza RUT y correos electrónicos por marcadores. Esto reduce "
    "la exposición innecesaria de datos personales dentro del historial."
)

add_title(doc, "8.5 Protección del frontend", 2)
add_bullets(
    doc,
    [
        "ProtectedRoute bloquea vistas internas sin sesión.",
        "El chatbot exige un token disponible antes de enviar consultas.",
        "El token se envía como Authorization: Bearer <token>.",
        "Una respuesta 401 elimina la sesión y solicita iniciar sesión nuevamente.",
        "OTP, biometría y acceso social simulados fueron deshabilitados para no generar tokens inválidos.",
    ],
)

add_page_break(doc)
add_title(doc, "9. Pruebas y validaciones realizadas", 1)
add_title(doc, "9.1 Validaciones automatizadas", 2)
add_table(
    doc,
    ["Validación", "Resultado"],
    [
        ["Backend: npm run build", "Correcto."],
        ["Backend: npm run test -- --runInBand", "9 de 9 pruebas aprobadas."],
        ["Frontend: npm run build", "Correcto."],
        ["Frontend: npm run lint", "Sin errores; permanece una advertencia previa en Alerts.tsx."],
        ["Prisma db push y generate", "Correcto; tabla chatbot_audit creada."],
    ],
    widths=[8, 8],
)

add_title(doc, "9.2 Casos probados en chatbot", 2)
add_bullets(
    doc,
    [
        "Rechazo de mensajes vacíos.",
        "Resumen de inventario.",
        "Separación entre productos sin stock y productos negativos.",
        "Orden correcto para menor stock.",
        "Orden correcto para mayor stock.",
        "Sanitización de RUT y correo antes de auditoría.",
        "Consulta sin token rechazada con HTTP 401.",
        "Login VIEWER genera JWT válido de tres segmentos.",
        "Consulta autenticada devuelve datos reales.",
        "Auditoría registra usuario, rol, consulta y tipo de respuesta.",
    ],
)

add_title(doc, "9.3 Datos reales observados durante validación", 2)
add_table(
    doc,
    ["Indicador", "Resultado observado"],
    [
        ["Productos registrados", "790"],
        ["Unidades disponibles normalizadas", "531.892"],
        ["Productos bajo mínimo", "107"],
        ["Productos sin stock", "1"],
        ["Productos que requieren ajuste", "22"],
        ["Producto con mayor stock observado", "CHICLE_FINI_*18, 24.858 unidades"],
    ],
    widths=[8, 8],
)

add_page_break(doc)
add_title(doc, "10. Ejecución local y dependencias", 1)
add_title(doc, "10.1 Requisitos generales", 2)
add_bullets(
    doc,
    [
        "Windows con PowerShell o una terminal equivalente.",
        "Node.js LTS y npm.",
        "Git para clonar y versionar repositorios.",
        "Visual Studio Code recomendado.",
        "No es necesario instalar SQLite por separado para ejecutar el prototipo.",
    ],
)

add_title(doc, "10.2 Ejecutar backend", 2)
add_code(
    doc,
    'cd "RUTA\\Pinval_SICD_back"\n'
    "npm install\n"
    "npx prisma generate\n"
    "npm run start:dev"
)
add_body(doc, "API esperada: http://localhost:3000")

add_title(doc, "10.3 Ejecutar frontend", 2)
add_code(
    doc,
    'cd "RUTA\\Pinval_SICD_front"\n'
    "npm install\n"
    "npm run dev"
)
add_body(doc, "Interfaz esperada: http://localhost:5173. Vite puede utilizar 5174 si el puerto está ocupado.")

add_title(doc, "10.4 Variables de entorno recomendadas", 2)
add_code(
    doc,
    'DATABASE_URL="file:./dev.db"\n'
    'JWT_SECRET="secreto-largo-y-privado"\n'
    'FRONTEND_URL="http://localhost:5173,http://localhost:5174"'
)

add_title(doc, "10.5 Usuarios de prueba", 2)
add_table(
    doc,
    ["Usuario", "Contraseña", "Rol"],
    [
        ["admin", "admin123", "ADMIN"],
        ["inventario", "stock123", "STOCK"],
        ["consulta", "consulta123", "VIEWER"],
    ],
    widths=[5, 5, 6],
)
add_callout(
    doc,
    "Advertencia",
    "Estas credenciales son únicamente de desarrollo. En una implementación real, los usuarios y contraseñas "
    "deben persistirse en base de datos, administrarse de forma segura y reemplazarse antes del despliegue.",
    fill="FFECEC",
)

add_page_break(doc)
add_title(doc, "11. Decisiones técnicas y justificación", 1)
add_table(
    doc,
    ["Decisión", "Justificación"],
    [
        ["Eliminar carga Excel del frontend", "La base de datos y el backend deben ser la fuente de verdad."],
        ["Usar SQLite inicialmente", "Simplifica el desarrollo y permite trabajar sin infraestructura adicional."],
        ["Mantener Prisma", "Reduce acoplamiento y facilita evolución hacia PostgreSQL."],
        ["Preservar stock negativo", "Es evidencia de una inconsistencia real que debe revisarse."],
        ["Separar sin stock de requiere ajuste", "Evita mezclar una condición operacional con un problema de datos."],
        ["Chatbot de solo lectura", "Reduce riesgos durante la primera integración."],
        ["No usar aún un modelo generativo", "Las consultas controladas son más predecibles, seguras y auditables."],
        ["Guardar sesión en sessionStorage", "Reduce persistencia del token frente a localStorage."],
        ["Proteger rutas y endpoint", "Evita acceso directo sin autenticación."],
        ["Auditoría sanitizada", "Mejora trazabilidad limitando exposición de datos personales."],
    ],
    widths=[6, 10],
)

add_title(doc, "12. Limitaciones, riesgos y trabajo pendiente", 1)
add_table(
    doc,
    ["Tema", "Estado actual", "Trabajo requerido"],
    [
        ["Usuarios", "Definidos en código con hashes bcrypt.", "Crear modelo User, administración y recuperación segura."],
        ["Ventas y compras", "Parte de la interfaz existe; datos transaccionales completos pendientes.", "Crear endpoints POST/GET y reglas atómicas de stock."],
        ["Integración POS", "No implementada.", "Definir mecanismo de sincronización y manejo de duplicados."],
        ["Stock mínimo", "Valor temporal fijo en backend.", "Persistir mínimo por producto."],
        ["Fechas y ganancias", "Información histórica insuficiente para cálculo preciso.", "Registrar cada compra y venta con fecha, costo y precio."],
        ["Migración PostgreSQL", "Preparada conceptualmente, no ejecutada.", "Diseñar migraciones y pruebas de concurrencia."],
        ["Chatbot con IA", "No incorporado.", "Evaluar modelo solo para clasificación de intención."],
        ["Modo bilingüe", "Infraestructura incorporada, cobertura puede ser parcial.", "Completar y revisar traducciones."],
        ["Pruebas frontend", "Build y lint.", "Agregar pruebas de componentes y flujos E2E."],
        ["Advertencia Alerts.tsx", "useMemo con dependencias faltantes.", "Corregir para evitar cálculo desactualizado."],
    ],
    widths=[4, 6, 6],
)

add_title(doc, "13. Próximos pasos recomendados", 1)
add_numbered(
    doc,
    [
        "Integrar cuidadosamente los cambios de las copias aisladas en ramas del repositorio oficial.",
        "Persistir usuarios, roles y contraseñas en base de datos.",
        "Crear endpoints transaccionales reales para compras, ventas y movimientos.",
        "Aplicar transacciones de base de datos para actualizar stock y registrar movimientos de forma atómica.",
        "Definir stock mínimo por producto y permitir su administración.",
        "Completar fechas de compras y ventas para calcular costos, ventas y márgenes por período.",
        "Crear una vista administrativa de auditoría.",
        "Agregar paginación, búsqueda y filtros del lado del backend.",
        "Diseñar pruebas E2E para login, roles, compras, ventas y chatbot.",
        "Evaluar posteriormente un modelo de IA limitado a interpretación de lenguaje, con disclaimer y controles.",
        "Planificar migración desde SQLite hacia PostgreSQL antes de operación multiusuario productiva.",
    ],
)

add_title(doc, "14. Conclusiones", 1)
add_body(
    doc,
    "El desarrollo permitió transformar una idea inicial centrada en visualizar columnas de Excel en una "
    "arquitectura web integrada con frontend, backend, base de datos, autenticación y consultas auditables. "
    "La interfaz evolucionó mediante revisiones sucesivas para reducir ruido visual y representar de forma más "
    "clara los procesos reales de Pinval."
)
add_body(
    doc,
    "El backend ya entrega productos normalizados y distingue correctamente entre disponibilidad, bajo mínimo, "
    "sin stock y registros negativos que requieren ajuste. El chatbot constituye una capa adicional de consulta "
    "que facilita el acceso a información sin entregar permisos de modificación."
)
add_body(
    doc,
    "La incorporación de JWT, roles, protección de rutas y auditoría representa un avance importante respecto del "
    "prototipo inicial. No obstante, para llegar a producción será necesario completar las operaciones "
    "transaccionales, persistir usuarios, fortalecer pruebas, integrar el POS y migrar a una base de datos adecuada "
    "para concurrencia."
)
add_callout(
    doc,
    "Resultado general",
    "SICD Pinval se encuentra en una etapa de prototipo funcional integrado, con una base técnica coherente para "
    "continuar desarrollando reglas de negocio, operaciones reales y controles de gobernanza de datos.",
    fill="E8F7EF",
)

# Document properties
doc.core_properties.title = "Informe de construcción y desarrollo del SICD Pinval"
doc.core_properties.subject = "Documentación técnica y funcional del proyecto SICD Pinval"
doc.core_properties.author = "Equipo SICD Pinval"
doc.core_properties.keywords = "SICD, Pinval, inventario, frontend, backend, chatbot, seguridad"

doc.save(OUTPUT)
print(OUTPUT.resolve())
